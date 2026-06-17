"""Word exporter — python-docx based formatted report."""

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..core.model import UnifiedProject
from .base import BaseExporter

HEADERS = [
    "#", "Clip Name", "Source File", "Source In", "Source Out",
    "Record In", "Record Out", "Duration", "Speed", "Transition",
]


class WordExporter(BaseExporter):
    @classmethod
    def format_name(cls) -> str:
        return "Microsoft Word (docx)"

    @classmethod
    def file_extension(cls) -> str:
        return ".docx"

    def export(self, project: UnifiedProject, output_path: Path) -> Path:
        doc = Document()

        # ── Styles ──
        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(6)

        # ── Title Page ──
        doc.add_heading("DEEPWORLD — Media Conversion Report", level=0)
        doc.add_paragraph(
            f"Source: {project.metadata.source_format.upper()}  |  "
            f"File: {project.metadata.source_path}"
        )
        doc.add_paragraph(
            f"Framerate: {float(project.metadata.framerate):.4f} fps  |  "
            f"Drop Frame: {project.metadata.drop_frame}"
        )
        doc.add_paragraph(f"Project: {project.metadata.title}")
        if project.metadata.author:
            doc.add_paragraph(f"Author: {project.metadata.author}")
        doc.add_paragraph("")

        total_clips = sum(
            len(c.clips) for t in project.timelines for c in t.tracks
        )

        # ── Executive Summary ──
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            f"This report contains {len(project.timelines)} timeline(s) "
            f"with a total of {total_clips} clips "
            f"across {sum(len(t.tracks) for t in project.timelines)} track(s)."
        )

        # ── Per Timeline Detail ──
        for ti, timeline in enumerate(project.timelines, 1):
            doc.add_heading(f"Timeline {ti}: {timeline.name}", level=2)

            for track in timeline.tracks:
                doc.add_heading(
                    f"Track {track.index}: {track.name} ({track.track_type.value})",
                    level=3,
                )

                if not track.clips:
                    doc.add_paragraph("No clips in this track.")
                    continue

                # Create table
                table = doc.add_table(rows=1 + len(track.clips), cols=len(HEADERS))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                for i, h in enumerate(HEADERS):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)

                # Data rows
                for j, clip in enumerate(track.clips):
                    row = table.rows[j + 1]
                    vals = [
                        str(j + 1),
                        clip.clip_name,
                        clip.source_file or "-",
                        clip.source_in.to_smpte_string() if clip.source_in else "-",
                        clip.source_out.to_smpte_string() if clip.source_out else "-",
                        clip.record_in.to_smpte_string() if clip.record_in else "-",
                        clip.record_out.to_smpte_string() if clip.record_out else "-",
                        clip.record_duration.to_smpte_string() if clip.record_duration else "-",
                        f"{clip.speed:.0f}%",
                        clip.transition.transition_type.value if clip.transition else "Cut",
                    ]
                    for k, v in enumerate(vals):
                        row.cells[k].text = v
                        for paragraph in row.cells[k].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)

                doc.add_paragraph("")  # spacing

        # ── Footer ──
        doc.add_paragraph("— End of Report —").alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(str(output_path))
        return output_path


# Register
from .base import ExporterRegistry
ExporterRegistry.register(WordExporter)
