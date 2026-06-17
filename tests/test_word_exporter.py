"""Tests for Word exporter."""

from pathlib import Path
import tempfile
from deepworld.parsers import ParserRegistry
from deepworld.exporters import ExporterRegistry
from docx import Document

SAMPLES = Path(__file__).parent.parent / "samples"
SAMPLE_EDL = SAMPLES / "sample_cmx3600.edl"


class TestWordExporter:
    def setup_method(self):
        self.parser = ParserRegistry.get_parser(SAMPLE_EDL)
        self.exporter = ExporterRegistry.get_exporter(Path("test.docx"))
        self.project = self.parser.parse(SAMPLE_EDL)

    def test_exporter_type(self):
        assert "Word" in self.exporter.format_name()

    def test_export_creates_file(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        try:
            result = self.exporter.export(self.project, path)
            assert result.exists()
            assert result.stat().st_size > 0
        finally:
            path.unlink(missing_ok=True)

    def test_export_has_content(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        try:
            self.exporter.export(self.project, path)
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs)
            assert "DEEPWORLD" in text
            assert "Timeline" in text or "Track" in text
        finally:
            path.unlink(missing_ok=True)

    def test_export_has_tables(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        try:
            self.exporter.export(self.project, path)
            doc = Document(str(path))
            assert len(doc.tables) > 0
        finally:
            path.unlink(missing_ok=True)
